"""
Document Parser Service
Parse .md, .txt, .pdf, .doc/.docx files with Chinese support
"""
import os
from pathlib import Path
from typing import Optional
from app.utils.logger import setup_logger

logger = setup_logger("document_parser")

# Maximum characters to extract from a document
MAX_DOCUMENT_CHARS = 50000

# Supported document extensions
DOCUMENT_EXTENSIONS = {'.md', '.txt', '.pdf', '.doc', '.docx'}


class DocumentParserService:
    """Service for parsing document files"""

    def is_document(self, file_path: str) -> bool:
        """Check if file is a parseable document"""
        if not file_path:
            return False
        ext = Path(file_path).suffix.lower()
        return ext in DOCUMENT_EXTENSIONS

    def parse_file(self, file_path: str) -> Optional[str]:
        """
        Parse document file and return text content

        Args:
            file_path: Path to the document file

        Returns:
            Extracted text content or None if parsing failed
        """
        if not file_path or not os.path.exists(file_path):
            logger.warning(f"File not found: {file_path}")
            return None

        ext = Path(file_path).suffix.lower()

        try:
            if ext in {'.md', '.txt'}:
                content = self._parse_text_file(file_path)
            elif ext == '.pdf':
                content = self._parse_pdf(file_path)
            elif ext in {'.doc', '.docx'}:
                content = self._parse_word(file_path)
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return None

            if content:
                # Truncate if too long
                if len(content) > MAX_DOCUMENT_CHARS:
                    content = content[:MAX_DOCUMENT_CHARS] + "\n\n[文档内容过长，已截断...]"
                    logger.info(f"Document truncated: {file_path}")

                logger.info(f"Successfully parsed {ext} file: {file_path}, length={len(content)}")
                return content
            else:
                logger.warning(f"No content extracted from: {file_path}")
                return None

        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            return None

    def _parse_text_file(self, file_path: str) -> Optional[str]:
        """
        Parse .md or .txt file with encoding detection
        Supports UTF-8, GBK, GB2312, and other encodings
        """
        try:
            # First try UTF-8
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                pass

            # Try to detect encoding using chardet
            try:
                import chardet
                with open(file_path, 'rb') as f:
                    raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected.get('encoding', 'utf-8')
                confidence = detected.get('confidence', 0)
                logger.info(f"Detected encoding: {encoding} (confidence: {confidence})")
                return raw_data.decode(encoding)
            except Exception as e:
                logger.warning(f"Chardet detection failed: {e}")

            # Fallback to common Chinese encodings
            for encoding in ['gb2312', 'gbk', 'gb18030', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue

            logger.error(f"Could not decode file with any encoding: {file_path}")
            return None

        except Exception as e:
            logger.error(f"Failed to parse text file {file_path}: {e}")
            return None

    def _parse_pdf(self, file_path: str) -> Optional[str]:
        """
        Parse PDF file using PyMuPDF (fitz)
        Better Chinese support than PyPDF2
        """
        try:
            import fitz  # PyMuPDF

            text_parts = []
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc, start=1):
                    text = page.get_text()
                    if text.strip():
                        text_parts.append(f"[第{page_num}页]\n{text.strip()}")

            if text_parts:
                return "\n\n".join(text_parts)
            else:
                logger.warning(f"No text extracted from PDF: {file_path}")
                return None

        except ImportError:
            logger.error("PyMuPDF not installed. Install with: pip install PyMuPDF")
            return None
        except Exception as e:
            logger.error(f"Failed to parse PDF {file_path}: {e}")
            return None

    def _parse_word(self, file_path: str) -> Optional[str]:
        """
        Parse Word document (.doc/.docx) using python-docx
        Note: .doc files need to be converted first, python-docx only supports .docx
        """
        try:
            from docx import Document

            ext = Path(file_path).suffix.lower()

            if ext == '.doc':
                # .doc files are not directly supported
                logger.warning(f".doc format not directly supported, only .docx: {file_path}")
                return "[不支持 .doc 格式，请转换为 .docx 格式后重新上传]"

            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_rows.append(row_text)
                if table_rows:
                    text_parts.append("\n[表格]\n" + "\n".join(table_rows))

            if text_parts:
                return "\n\n".join(text_parts)
            else:
                logger.warning(f"No text extracted from Word document: {file_path}")
                return None

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Failed to parse Word document {file_path}: {e}")
            return None


# Singleton instance
document_parser = DocumentParserService()
